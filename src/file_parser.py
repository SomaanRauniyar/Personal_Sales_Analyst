import pandas as pd
import pdfplumber
from docx import Document
import os
from io import BytesIO

# -----------------------------
# CSV Parsing
# -----------------------------
def parse_csv(file_obj):
    """
    Reads a CSV/TSV file (path or file-like object) and returns all rows as a list of dicts.
    Automatically detects delimiters (comma, tab, semicolon) using pandas' parser.
    """
    def _read(obj):
        # Try utf-8 with auto delimiter detection; fall back to latin-1
        try:
            return pd.read_csv(obj, sep=None, engine='python', encoding='utf-8')
        except Exception:
            if isinstance(obj, BytesIO):
                obj.seek(0)
            return pd.read_csv(obj, sep=None, engine='python', encoding='latin-1')

    if isinstance(file_obj, BytesIO):
        file_obj.seek(0)
        df = _read(file_obj)
    else:
        df = _read(file_obj)

    # Basic cleanup: strip stray quotes/whitespace and coerce numerics
    for col in df.columns:
        if df[col].dtype == object:
            df[col] = df[col].astype(str).str.strip().str.strip('"\'')
            # Try numeric coercion where possible
            coerced = pd.to_numeric(df[col], errors='ignore')
            df[col] = coerced

    return df.to_dict(orient='records')


# -----------------------------
# PDF Parsing
# -----------------------------
def parse_pdf(file_obj):
    """
    Reads a PDF and attempts to extract tabular data.
    Returns:
      - if tables detected: list of dict rows
      - else: list of text paragraphs
    """
    def _extract(pdf):
        rows = []
        for page in pdf.pages:
            # try tables first
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            for tbl in tables:
                if not tbl or len(tbl) < 2:
                    continue
                header = tbl[0]
                for r in tbl[1:]:
                    rec = {}
                    for i, h in enumerate(header):
                        key = str(h).strip() if h is not None else f"col_{i}"
                        val = r[i] if i < len(r) else None
                        rec[key] = val
                    rows.append(rec)
        if rows:
            return rows
        # fallback: plain text paragraphs
        texts = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return texts

    if isinstance(file_obj, BytesIO):
        file_obj.seek(0)
        with pdfplumber.open(file_obj) as pdf:
            return _extract(pdf)
    else:
        with pdfplumber.open(file_obj) as pdf:
            return _extract(pdf)


# -----------------------------
# Word (DOCX) Parsing
# -----------------------------
def parse_docx(file_obj):
    """
    Reads a Word document and tries to extract tables first.
    Returns list of dict rows if tables found, else list of paragraphs.
    """
    if isinstance(file_obj, BytesIO):
        file_obj.seek(0)
        doc = Document(file_obj)
    else:
        doc = Document(file_obj)

    # try tables
    rows = []
    try:
        for t in doc.tables:
            if len(t.rows) < 2:
                continue
            header = [c.text.strip() or f"col_{i}" for i, c in enumerate(t.rows[0].cells)]
            for r in t.rows[1:]:
                rec = {}
                for i, c in enumerate(r.cells):
                    key = header[i] if i < len(header) else f"col_{i}"
                    rec[key] = c.text
                rows.append(rec)
    except Exception:
        pass
    if rows:
        return rows

    # fallback paragraphs
    texts = [para.text for para in doc.paragraphs if para.text.strip() != '']
    return texts


# -----------------------------
# General Parser
# -----------------------------
def parse_file(file_obj, filename=None):
    """
    Detects file type from filename (if provided) and parses accordingly.
    Returns list of text blocks.
    """
    if filename is None:
        raise ValueError("filename must be provided to detect file type")

    _, ext = os.path.splitext(filename.lower())
    if ext == '.csv':
        # Return structured rows (list[dict]) for downstream DataFrame previewing
        # Stringification can be done by callers when needed for embeddings
        data = parse_csv(file_obj)
        return data
    elif ext == '.pdf':
        return parse_pdf(file_obj)
    elif ext in ['.docx', '.doc']:
        return parse_docx(file_obj)
    else:
        raise ValueError(f"Unsupported file type: {ext}")